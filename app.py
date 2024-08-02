from fastapi import FastAPI, File, UploadFile, Form
from fastapi.responses import FileResponse
from typing import Dict
from docx import Document
import tempfile
import json

app = FastAPI()

def substituir_variaveis(doc, substituicoes):
    """
    Substitui variáveis no conteúdo de um documento DOCX.
    """
    for paragrafo in doc.paragraphs:
        for run in paragrafo.runs:
            for chave, valor in substituicoes.items():
                if chave in run.text:
                    run.text = run.text.replace(chave, valor)

    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    for run in paragrafo.runs:
                        for chave, valor in substituicoes.items():
                            if chave in run.text:
                                run.text = run.text.replace(chave, valor)


@app.post("/processar_docx/")
async def processar_docx(file: UploadFile = File(...), substituicoes: str = Form(...)):
    # Parse the substituicoes JSON string
    substituicoes_dict = json.loads(substituicoes)

    # Ler o documento DOCX
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp.write(file.file.read())
        tmp_path = tmp.name

    doc = Document(tmp_path)

    # Substituir variáveis
    substituir_variaveis(doc, substituicoes_dict["substituicoes"])

    # Salvar o documento modificado
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_out:
        doc.save(tmp_out.name)
        tmp_out_path = tmp_out.name
        
    # Retornar o arquivo DOCX modificado
    return FileResponse(tmp_out_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename="resultado.docx")

if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)
